#!/usr/bin/env python3
"""speech_to_docx.py — Convert a speech Markdown file to a styled DOCX.

Usage:
    python speech_to_docx.py input.md [output.docx]

Requires: python-docx
"""

import sys
import re
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("Install python-docx:  pip install python-docx")
    sys.exit(1)


def md_to_docx(md_path: str, docx_path: str | None = None):
    md = Path(md_path)
    if not md.exists():
        print(f"File not found: {md_path}")
        sys.exit(1)

    out = Path(docx_path) if docx_path else md.with_suffix(".docx")
    text = md.read_text(encoding="utf-8")

    doc = Document()

    # -- styles ----------------------------------------------------------
    style_normal = doc.styles["Normal"]
    style_normal.font.name = "Calibri"
    style_normal.font.size = Pt(12)

    for line in text.splitlines():
        line = line.rstrip()

        # headings
        if line.startswith("# "):
            p = doc.add_heading(line[2:], level=1)
        elif line.startswith("## "):
            p = doc.add_heading(line[3:], level=2)
        elif line.startswith("### "):
            p = doc.add_heading(line[4:], level=3)
        # stage directions  [PAUSE] / [SLIDE]
        elif re.match(r"^\[.*\]$", line.strip()):
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.italic = True
            run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
        # blank line → small spacer
        elif line == "":
            doc.add_paragraph("")
        else:
            # bold fragments
            p = doc.add_paragraph()
            parts = re.split(r"(\*\*.*?\*\*)", line)
            for part in parts:
                if part.startswith("**") and part.endswith("**"):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    p.add_run(part)

    doc.save(str(out))
    print(f"Saved → {out}")


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print(__doc__)
        sys.exit(0)
    md_to_docx(sys.argv[1], sys.argv[2] if len(sys.argv) > 2 else None)